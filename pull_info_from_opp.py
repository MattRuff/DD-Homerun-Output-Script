#!/usr/bin/env python3
from __future__ import annotations
"""
Homerun Presales data exporter.

Pulls evaluation data (general info, meeting minutes, transcripts) from the
Homerun API and exports it as JSON or plain text for LLM post-processing.

Auth is automatic — reads cookies from your Chrome session via rookiepy.
Just be logged in to Homerun in Chrome.

Usage:
    python pull_info_from_opp.py --list                     # list your opps
    python pull_info_from_opp.py --all                      # export all opps
    python pull_info_from_opp.py <uuid_or_name>             # export one opp
    python pull_info_from_opp.py --list --team 'Name1' 'Name2'   # team view
"""

import argparse
import base64
import concurrent.futures
import json
import os
import re
import sys
import time
import warnings
warnings.filterwarnings("ignore", message="urllib3.*doesn't match a supported version")
from html.parser import HTMLParser
from urllib.parse import quote

import requests

BASE_URL = os.environ.get(
    "HOMERUN_BASE_URL",
    "https://datadog.cloud.homerunpresales.com/api/v1",
)


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

def _check_jwt_expiry(cookie_str: str, source: str) -> None:
    """Warn and exit early if the jwttoken in the cookie string is expired."""
    for part in cookie_str.split(";"):
        part = part.strip()
        if part.startswith("jwttoken="):
            token = part[len("jwttoken="):]
            try:
                payload_b64 = token.split(".")[1]
                payload_b64 += "=" * (-len(payload_b64) % 4)
                payload = json.loads(base64.urlsafe_b64decode(payload_b64))
                exp = payload.get("exp")
                if exp and exp < time.time():
                    hours_ago = (time.time() - exp) / 3600
                    if hours_ago > 1:
                        print(
                            f"Error: JWT token from {source} expired {hours_ago:.0f}h ago.\n"
                            f"  Fix: open Homerun in Chrome, do a hard refresh (Cmd+Shift+R),\n"
                            f"  wait a few seconds for the page to fully load, then re-run.\n"
                            f"  (Homerun issues short-lived tokens; the page must fully load\n"
                            f"  so the refresh-token exchange writes a new cookie to Chrome.)",
                            file=sys.stderr,
                        )
                        sys.exit(1)
                    else:
                        print(
                            f"Warning: JWT token from {source} recently expired. "
                            f"Attempting anyway (refresh token may still work)...",
                            file=sys.stderr,
                        )
            except (IndexError, ValueError, json.JSONDecodeError):
                pass
            break


def _get_cookies(args) -> str:
    """Resolve cookies: explicit arg > file > env > Chrome cookie store."""
    if args.cookies:
        _check_jwt_expiry(args.cookies, "--cookies flag")
        return args.cookies
    if args.cookies_file:
        with open(args.cookies_file, encoding="utf-8") as f:
            cookie_str = f.read().strip()
        _check_jwt_expiry(cookie_str, f"file {args.cookies_file}")
        return cookie_str
    if os.environ.get("HOMERUN_COOKIES"):
        cookie_str = os.environ["HOMERUN_COOKIES"]
        _check_jwt_expiry(cookie_str, "HOMERUN_COOKIES env var")
        return cookie_str

    try:
        import rookiepy
    except ImportError:
        print(
            "Error: no cookies provided and rookiepy is not installed.\n"
            "  Install rookiepy (pip install rookiepy) and log in to Homerun in Chrome,\n"
            "  or pass cookies explicitly via --cookies, -f, or HOMERUN_COOKIES env var.",
            file=sys.stderr,
        )
        sys.exit(1)

    try:
        cookies = rookiepy.chrome(["homerunpresales.com"])
    except RuntimeError as e:
        err = str(e).lower()
        if "can't find cookies" in err or "no such file" in err:
            if sys.platform == "darwin":
                print(
                    "Error: could not read Chrome cookie database.\n"
                    "  On macOS this usually means your terminal app lacks Full Disk Access.\n"
                    "  Fix: System Settings > Privacy & Security > Full Disk Access\n"
                    "        -> toggle ON your terminal (Terminal, iTerm2, Cursor, VS Code, etc.)\n"
                    "        -> restart the terminal and try again.\n"
                    "  If running in Docker, pass cookies via -e HOMERUN_COOKIES or --cookies.",
                    file=sys.stderr,
                )
            else:
                print(
                    "Error: could not read Chrome cookies (cookie store not found).\n"
                    "  In Docker / headless, supply cookies with one of:\n"
                    "    docker run -e HOMERUN_COOKIES='jwttoken=...; jrtoken=...' IMAGE --all\n"
                    "    docker run IMAGE --cookies 'jwttoken=...; jrtoken=...' --all",
                    file=sys.stderr,
                )
        else:
            print(
                f"Error: rookiepy could not read Chrome cookies: {e}\n"
                "  On macOS, check:\n"
                "    1. Full Disk Access granted to your terminal app\n"
                "    2. Keychain prompt: click 'Always Allow' for 'Chrome Safe Storage'\n"
                "  Or pass cookies via --cookies, -f, or HOMERUN_COOKIES env var.",
                file=sys.stderr,
            )
        sys.exit(1)
    except OSError as e:
        print(
            f"Error: OS permission error reading Chrome cookies: {e}\n"
            "  On macOS: grant Full Disk Access to your terminal app:\n"
            "    System Settings > Privacy & Security > Full Disk Access\n"
            "  Then restart the terminal and try again.",
            file=sys.stderr,
        )
        sys.exit(1)

    if not cookies:
        print("No Homerun cookies in Chrome. Log in first.", file=sys.stderr)
        sys.exit(1)

    if not any(c["name"] == "jwttoken" for c in cookies):
        print(
            "Error: jwttoken not found in Chrome cookies.\n"
            "  Log in to Homerun in Chrome and try again.",
            file=sys.stderr,
        )
        sys.exit(1)

    seen = {}
    for c in cookies:
        seen[c["name"]] = c["value"]
    cookie_str = "; ".join(f"{k}={v}" for k, v in seen.items())
    _check_jwt_expiry(cookie_str, "Chrome cookies")
    return cookie_str


def _headers(cookies: str) -> dict:
    return {
        "accept": "application/json",
        "referer": BASE_URL.rsplit("/api/", 1)[0] + "/",
        "user-agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/131.0.0.0 Safari/537.36",
        "Cookie": cookies,
    }


# ---------------------------------------------------------------------------
# API helpers
# ---------------------------------------------------------------------------

def _raise_on_auth_error(r: requests.Response) -> None:
    """Check for auth-related failures and give a clear message before exiting."""
    if r.status_code in (401, 403, 500):
        body = ""
        try:
            body = r.text[:200]
        except Exception:
            pass
        if "expired" in body.lower() or "token" in body.lower() or r.status_code == 401:
            print(
                f"Error: API returned {r.status_code} — likely expired or invalid session.\n"
                f"  Response: {body}\n"
                f"  Fix: log in to Homerun in Chrome and re-run (rookiepy grabs fresh cookies).",
                file=sys.stderr,
            )
            sys.exit(1)
    r.raise_for_status()


def _get_csrf(cookies: str) -> str:
    """Fetch a CSRF token (required for POST requests)."""
    r = requests.get(f"{BASE_URL}/authenticate/csrf", headers=_headers(cookies))
    _raise_on_auth_error(r)
    return r.json()["xsrf_token"]


def _get(cookies: str, path: str) -> dict | list:
    r = requests.get(f"{BASE_URL}/{path}", headers=_headers(cookies))
    _raise_on_auth_error(r)
    return r.json()


def _post(cookies: str, path: str, payload: dict) -> dict:
    hdrs = _headers(cookies)
    hdrs["content-type"] = "application/json"
    hdrs["x-csrf-token"] = _get_csrf(cookies)
    r = requests.post(f"{BASE_URL}/{path}", headers=hdrs, json=payload)
    _raise_on_auth_error(r)
    return r.json()


# ---------------------------------------------------------------------------
# Data fetching
# ---------------------------------------------------------------------------

# Salesforce stage filter values used by fetch_opportunities(stage_filter="active").
# These are standard Homerun stage names; override via SALESFORCE_STAGES_1_5 if
# your instance uses different values.
SALESFORCE_STAGES_1_5 = [
    {"value_uuid": "StageName_1. Conducting Discovery", "display_value": "1. Conducting Discovery"},
    {"value_uuid": "StageName_2. Demonstrating & Champion Building", "display_value": "2. Demonstrating & Champion Building"},
    {"value_uuid": "StageName_3. Proving Value", "display_value": "3. Proving Value"},
    {"value_uuid": "StageName_4. Defining Deal Structure & Proposal", "display_value": "4. Defining Deal Structure & Proposal"},
    {"value_uuid": "StageName_5. Finalizing Paper Process", "display_value": "5. Finalizing Paper Process"},
]

# Optional local cache of user display-name → UUID for --team lookups.
# Avoids an API call when names are known. The script falls back to the
# /user endpoint for any name not listed here. Populate with your own
# team members or leave empty.
KNOWN_USERS: dict[str, dict] = {
    # "jane doe": {"value_uuid": "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx", "display_value": "Jane Doe"},
}


def _resolve_users(cookies: str, names: list[str]) -> list[dict]:
    """Resolve user names to UUID filter values. Uses local cache, falls back to API."""
    matches, need_lookup = [], []
    for n in names:
        cached = KNOWN_USERS.get(n.lower())
        if cached:
            matches.append(cached)
        else:
            need_lookup.append(n)

    if need_lookup:
        try:
            all_users = _get(cookies, "user")
            lookup_lower = [n.lower() for n in need_lookup]
            for u in all_users:
                full = u.get("full_name", "").lower()
                if full in lookup_lower:
                    matches.append({"value_uuid": u["uuid"], "display_value": u["full_name"]})
        except requests.HTTPError:
            pass

    found = {m["display_value"].lower() for m in matches}
    missing = [n for n in names if n.lower() not in found]
    if missing:
        print(f"Warning: could not find users: {', '.join(missing)}", file=sys.stderr)
    return matches


def fetch_opportunities(cookies: str, user_values: list[dict] | None = None,
                        stage_filter: str = "undecided") -> list[dict]:
    """
    Fetch opportunities filtered by tech leads and stage.

    user_values: list of {"value_uuid": ..., "display_value": ...} for Tech Lead filter.
                 If None, uses the authenticated user.
    stage_filter: "undecided" (Deal Win = Undecided) or "active" (stages 1-5).
    """
    if user_values is None:
        user = _get(cookies, "authenticated")
        user_values = [{"value_uuid": user["uuid"], "display_value": user.get("full_name", "")}]

    data_filters = [
        {
            "analytic_field_uuid": "0",
            "display_name": "Technical Lead",
            "filter_field_type": "userse",
            "field_type": "value",
            "fact_table": "user",
            "user_attribute_uuid": "",
            "comparator_type": "is_any",
            "uuid": "",
            "values": [{"uuid": "", **v} for v in user_values],
        },
    ]

    if stage_filter == "active":
        data_filters.append({
            "analytic_field_uuid": "1cc38e7a-83b4-437c-8689-8e427a993c8d",
            "display_name": "Salesforce Stage",
            "filter_field_type": "user_attribute_value",
            "field_type": "value",
            "fact_table": "b43ad0b8f5ac4d7692051b6d2af734d2_attr_value",
            "user_attribute_uuid": "b43ad0b8-f5ac-4d76-9205-1b6d2af734d2",
            "comparator_type": "is_any",
            "uuid": "",
            "values": [{"uuid": "", **v} for v in SALESFORCE_STAGES_1_5],
        })
    else:
        data_filters.append({
            "analytic_field_uuid": "10",
            "display_name": "Deal Win",
            "filter_field_type": "evaluation_deal_win",
            "field_type": "value",
            "fact_table": "evaluation",
            "user_attribute_uuid": "",
            "comparator_type": "is_any",
            "uuid": "",
            "values": [{"uuid": "", "value_uuid": "Undecided", "display_value": "Undecided"}],
        })

    payload = {
        "chart_type": "Table",
        "collaborator_comparator": "or",
        "data_fields": [
            {"field_uuid": "2", "field_type": "value"},    # Opportunity Name
            {"field_uuid": "13", "field_type": "number"},   # Deal Size
            {"field_uuid": "4", "field_type": "value"},     # Homerun Stage
            {"field_uuid": "22", "field_type": "value"},    # Opportunity UUID
            {"field_uuid": "0", "field_type": "value"},     # Technical Lead
        ],
        "data_filters": data_filters,
        "order_by": "13",
        "order_by_direction": "Descending",
        "max_rows": 200,
    }
    resp = _post(cookies, "analytics/data", payload)
    labels = resp["labels"]
    idx = {name: labels.index(name) for name in
           ["Opportunity Name", "Opportunity UUID", "Homerun Stage", "Deal Size", "Technical Lead"]}
    return [
        {"name": row[idx["Opportunity Name"]], "uuid": row[idx["Opportunity UUID"]],
         "stage": row[idx["Homerun Stage"]], "deal_size": row[idx["Deal Size"]],
         "tech_lead": row[idx["Technical Lead"]]}
        for row in resp["series"][0]["data"]
    ]


def _fetch_single_transcript(cookies: str, uuid: str, item: dict, idx: int,
                              debug: bool = False) -> dict | None:
    """Fetch one transcript by its source_id. Returns None if skipped."""
    if not isinstance(item, dict):
        return None
    call_id = item.get("source_id") or item.get("callID") or item.get("call_id")
    name = item.get("meeting_name") or item.get("name")
    if not call_id:
        return None

    qs = f"callID={quote(call_id, safe='')}&evaluationUUID={quote(uuid, safe='')}"
    data = _get(cookies, f"transcript?{qs}")

    if debug:
        size = len(data) if isinstance(data, list) else "dict"
        print(f"DEBUG transcript {idx+1}: {name} -> {size}", file=sys.stderr)

    if isinstance(data, list):
        data = {"meeting_name": name, "segments": data}
    elif name:
        data["meeting_name"] = name
    return data


def fetch_evaluation_data(uuid: str, cookies: str, debug: bool = False) -> dict:
    """Fetch general info, meeting minutes, and transcripts for one evaluation."""
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as pool:
        f_info = pool.submit(_get, cookies, f"evaluation/{uuid}")
        f_minutes = pool.submit(_get, cookies, f"evaluation/{uuid}/meetingminutes")
        f_tlist = pool.submit(_get, cookies, f"evaluation/{uuid}/transcripts")

    results = {
        "general_info": f_info.result(),
        "meeting_minutes": f_minutes.result(),
    }

    transcripts_list = f_tlist.result()
    if debug:
        print(f"DEBUG transcripts_list: {len(transcripts_list)} items", file=sys.stderr)

    items = [t for t in (transcripts_list if isinstance(transcripts_list, list) else [])
             if isinstance(t, dict)]

    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as pool:
        futures = {
            pool.submit(_fetch_single_transcript, cookies, uuid, item, i, debug): i
            for i, item in enumerate(items)
        }
        transcript_results = [None] * len(items)
        for future in concurrent.futures.as_completed(futures):
            idx = futures[future]
            transcript_results[idx] = future.result()

    results["transcripts"] = [t for t in transcript_results if t is not None]
    return results


# ---------------------------------------------------------------------------
# Output formatting
# ---------------------------------------------------------------------------

class _HTMLStripper(HTMLParser):
    """Lightweight HTML-to-text converter."""
    def __init__(self):
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data):
        self._parts.append(data)

    def get_text(self) -> str:
        return "".join(self._parts)


def _strip_html(html: str) -> str:
    s = _HTMLStripper()
    s.feed(html)
    return s.get_text().strip()


_STATUS_RE = re.compile(
    r"(?:^|\n)\s*([A-Z]{2,4})\s+(\d{1,2}/\d{1,2}/\d{2,4})\s*[-\u2013\u2014]\s*(.*?)(?=\n\s*[A-Z]{2,4}\s+\d{1,2}/|\Z)",
    re.DOTALL,
)


def _parse_current_status(text: str) -> list[dict]:
    entries = []
    for m in _STATUS_RE.finditer(text):
        author, raw_date, note = m.group(1), m.group(2), m.group(3).strip()
        parts = raw_date.split("/")
        if len(parts) == 3:
            mm, dd, yy = parts
            yyyy = f"20{yy}" if len(yy) == 2 else yy
            date_str = f"{yyyy}-{int(mm):02d}-{int(dd):02d}"
        else:
            date_str = raw_date
        entries.append({"date": date_str, "author": author, "note": note})
    return entries


def _parse_opp_name(full_name: str) -> dict:
    """Parse 'Customer - Type - Term - Year' into components."""
    parts = [p.strip() for p in full_name.split(" - ", 3)]
    result = {"Customer": parts[0] if parts else full_name}
    if len(parts) >= 2:
        result["Type"] = parts[1]
    if len(parts) >= 3:
        result["Term"] = parts[2]
    if len(parts) >= 4:
        year_match = re.match(r"(\d{4})", parts[3])
        result["Year"] = int(year_match.group(1)) if year_match else parts[3]
        suffix = parts[3][year_match.end():].strip() if year_match else ""
        if suffix:
            result["Suffix"] = suffix
    return result


def _latest_date(*dates: str) -> str:
    valid = [d for d in dates if d and not d.startswith("0001")]
    return max(valid)[:10] if valid else ""


def build_opportunity_json(data: dict) -> dict:
    """Transform raw API data into the structured opportunity schema."""
    info = data["general_info"]
    name_parts = _parse_opp_name(info.get("eval_customer", ""))

    opp = {
        "Opportunity_ID": info.get("sfc_opportunity_id", ""),
        "Opportunity_Name": info.get("eval_customer", ""),
        **name_parts,
        "Stage": info.get("eval_status", ""),
        "Deal_Size": float(info.get("deal_size", 0)),
        "SE": info.get("eval_se_owner_full_name", ""),
        "SR": info.get("eval_sr_owner_full_name", ""),
        "Created_Date": info.get("eval_created_date", ""),
        "Last_Updated": _latest_date(
            info.get("last_updated_by_user", ""),
            info.get("last_updated_by_crm", ""),
            info.get("last_updated_by_workflow", ""),
        ),
        "Sentiment": info.get("eval_sentiment", ""),
        "Tech_Win": info.get("eval_tech_win", ""),
        "Deal_Win": info.get("eval_deal_win", ""),
        "Current_Status": _parse_current_status(info.get("eval_current_status", "")),
    }

    minutes = data.get("meeting_minutes", [])
    if isinstance(minutes, list):
        opp["Meeting_Minutes"] = [
            {
                "date": (m.get("mm_date") or "")[:10],
                "title": m.get("mm_title", ""),
                "summary": _strip_html(m.get("mm_content", "")),
            }
            for m in minutes
        ]
    else:
        opp["Meeting_Minutes"] = []

    transcripts_out = []
    for t in data.get("transcripts", []):
        if not isinstance(t, dict):
            continue
        segments = []
        for seg in t.get("segments", []):
            if not isinstance(seg, dict):
                continue
            segments.append({
                "speaker": seg.get("metadata_speaker", ""),
                "timestamp": seg.get("metadata_timestamp", ""),
                "text": seg.get("text", ""),
            })
        if segments:
            transcripts_out.append({
                "meeting": t.get("meeting_name", ""),
                "segments": segments,
            })
    opp["Transcripts"] = transcripts_out

    return opp


def build_output_envelope(opportunities: list[dict]) -> dict:
    return {
        "primary_key": "Opportunity_ID",
        "total_expected_opportunities": len(opportunities),
        "opportunities": opportunities,
    }


def format_opportunity_text(opp: dict) -> str:
    """Format a single opportunity as human-readable plain text."""
    lines = []
    lines.append(f"{'='*72}")
    lines.append(f"  {opp.get('Opportunity_Name', 'Unknown')}")
    lines.append(f"{'='*72}")
    lines.append("")
    lines.append(f"  Opportunity ID : {opp.get('Opportunity_ID', '')}")
    lines.append(f"  Customer       : {opp.get('Customer', '')}")
    if opp.get("Type"):
        lines.append(f"  Type           : {opp['Type']}")
    if opp.get("Term"):
        lines.append(f"  Term           : {opp['Term']}")
    if opp.get("Year"):
        lines.append(f"  Year           : {opp['Year']}")
    lines.append(f"  Stage          : {opp.get('Stage', '')}")
    lines.append(f"  Deal Size      : ${opp.get('Deal_Size', 0):,.2f}")
    lines.append(f"  SE             : {opp.get('SE', '')}")
    lines.append(f"  SR             : {opp.get('SR', '')}")
    lines.append(f"  Created        : {opp.get('Created_Date', '')}")
    lines.append(f"  Last Updated   : {opp.get('Last_Updated', '')}")
    lines.append(f"  Sentiment      : {opp.get('Sentiment', '')}")
    lines.append(f"  Tech Win       : {opp.get('Tech_Win', '')}")
    lines.append(f"  Deal Win       : {opp.get('Deal_Win', '')}")

    status_entries = opp.get("Current_Status", [])
    if status_entries:
        lines.append("")
        lines.append("  --- Current Status ---")
        for s in status_entries:
            lines.append(f"  [{s.get('date', '')}] {s.get('author', '')}: {s.get('note', '')}")

    minutes = opp.get("Meeting_Minutes", [])
    if minutes:
        lines.append("")
        lines.append(f"  --- Meeting Minutes ({len(minutes)}) ---")
        for m in minutes:
            lines.append(f"  [{m.get('date', '')}] {m.get('title', '')}")
            summary = m.get("summary", "")
            if summary:
                for para in summary.split("\n"):
                    para = para.strip()
                    if para:
                        lines.append(f"    {para}")
            lines.append("")

    transcripts = opp.get("Transcripts", [])
    if transcripts:
        total_segs = sum(len(m.get("segments", [])) for m in transcripts)
        lines.append(f"  --- Transcripts ({len(transcripts)} meetings, {total_segs} segments) ---")
        for meeting in transcripts:
            lines.append(f"  [{meeting.get('meeting', '')}]")
            for seg in meeting.get("segments", []):
                speaker = seg.get("speaker", "")
                text = seg.get("text", "")
                lines.append(f"    {speaker}: {text}")
            lines.append("")

    return "\n".join(lines)


def format_opportunity_docx(opp: dict) -> "docx.Document":
    """Format a single opportunity as a Word (.docx) document."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(opp.get("Opportunity_Name", "Unknown"), level=1)

    info_fields = [
        ("Opportunity ID", opp.get("Opportunity_ID", "")),
        ("Customer", opp.get("Customer", "")),
    ]
    if opp.get("Type"):
        info_fields.append(("Type", opp["Type"]))
    if opp.get("Term"):
        info_fields.append(("Term", opp["Term"]))
    if opp.get("Year"):
        info_fields.append(("Year", opp["Year"]))
    info_fields += [
        ("Stage", opp.get("Stage", "")),
        ("Deal Size", f"${opp.get('Deal_Size', 0):,.2f}"),
        ("SE", opp.get("SE", "")),
        ("SR", opp.get("SR", "")),
        ("Created", opp.get("Created_Date", "")),
        ("Last Updated", opp.get("Last_Updated", "")),
        ("Sentiment", opp.get("Sentiment", "")),
        ("Tech Win", opp.get("Tech_Win", "")),
        ("Deal Win", opp.get("Deal_Win", "")),
    ]

    table = doc.add_table(rows=len(info_fields), cols=2, style="Light Grid Accent 1")
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.2)
    for i, (label, value) in enumerate(info_fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.style.font.size = Pt(10)

    status_entries = opp.get("Current_Status", [])
    if status_entries:
        doc.add_heading("Current Status", level=2)
        for s in status_entries:
            p = doc.add_paragraph()
            run = p.add_run(f"[{s.get('date', '')}] {s.get('author', '')}: ")
            run.bold = True
            p.add_run(s.get("note", ""))

    minutes = opp.get("Meeting_Minutes", [])
    if minutes:
        doc.add_heading(f"Meeting Minutes ({len(minutes)})", level=2)
        for m in minutes:
            doc.add_heading(f"{m.get('date', '')} — {m.get('title', '')}", level=3)
            summary = m.get("summary", "")
            if summary:
                for para in summary.split("\n"):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)

    transcripts = opp.get("Transcripts", [])
    if transcripts:
        total_segs = sum(len(mt.get("segments", [])) for mt in transcripts)
        doc.add_heading(f"Transcripts ({len(transcripts)} meetings, {total_segs} segments)", level=2)
        for meeting in transcripts:
            doc.add_heading(meeting.get("meeting", ""), level=3)
            for seg in meeting.get("segments", []):
                p = doc.add_paragraph()
                run = p.add_run(f"{seg.get('speaker', '')}: ")
                run.bold = True
                p.add_run(seg.get("text", ""))

    return doc


def format_all_opportunities_text(opportunities: list[dict]) -> str:
    parts = [f"Total opportunities: {len(opportunities)}\n"]
    for opp in opportunities:
        parts.append(format_opportunity_text(opp))
    return "\n".join(parts)


def _safe_filename(name: str, max_len: int = 60) -> str:
    return name.replace("/", "-").replace("\\", "-").replace(":", "-")[:max_len]


_UUID_RE = re.compile(
    r"^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}$"
)


def _search_all_opps(cookies: str) -> list[dict]:
    """Search across ALL undecided opps (any tech lead) with a high row limit."""
    payload = {
        "chart_type": "Table",
        "collaborator_comparator": "or",
        "data_fields": [
            {"field_uuid": "2", "field_type": "value"},
            {"field_uuid": "22", "field_type": "value"},
            {"field_uuid": "0", "field_type": "value"},
        ],
        "data_filters": [
            {
                "analytic_field_uuid": "10",
                "display_name": "Deal Win",
                "filter_field_type": "evaluation_deal_win",
                "field_type": "value",
                "fact_table": "evaluation",
                "user_attribute_uuid": "",
                "comparator_type": "is_any",
                "uuid": "",
                "values": [{"uuid": "", "value_uuid": "Undecided", "display_value": "Undecided"}],
            },
        ],
        "order_by": "2",
        "order_by_direction": "Ascending",
        "max_rows": 5000,
    }
    resp = _post(cookies, "analytics/data", payload)
    labels = resp["labels"]
    name_idx = labels.index("Opportunity Name")
    uuid_idx = labels.index("Opportunity UUID")
    lead_idx = labels.index("Technical Lead")
    return [
        {"name": row[name_idx], "uuid": row[uuid_idx], "tech_lead": row[lead_idx]}
        for row in resp["series"][0]["data"]
    ]


def _resolve_opp_identifier(cookies: str, identifier: str) -> str:
    """Return evaluation UUID. Accepts a UUID or an opportunity name (exact or partial match)."""
    if _UUID_RE.match(identifier.strip()):
        return identifier.strip()

    name_lower = identifier.strip().lower()

    # First try: search the user's own opps (fast)
    try:
        opps = fetch_opportunities(cookies, stage_filter="active")
        for opp in opps:
            if opp.get("name", "").strip().lower() == name_lower:
                return opp["uuid"]
        for opp in opps:
            if name_lower in opp.get("name", "").lower():
                return opp["uuid"]
    except (requests.HTTPError, SystemExit):
        pass

    # Fallback: search ALL opps across all tech leads
    print("Not found in your opps, searching all opportunities...", file=sys.stderr)
    all_opps = _search_all_opps(cookies)
    for opp in all_opps:
        if opp.get("name", "").strip().lower() == name_lower:
            print(f"Found: {opp['name']} (Tech Lead: {opp['tech_lead']})", file=sys.stderr)
            return opp["uuid"]
    for opp in all_opps:
        if name_lower in opp.get("name", "").lower():
            print(f"Found: {opp['name']} (Tech Lead: {opp['tech_lead']})", file=sys.stderr)
            return opp["uuid"]

    print(f"No opportunity found matching: {identifier!r}", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Homerun Presales data exporter")
    parser.add_argument("evaluation_uuid", nargs="?",
                        help="Evaluation UUID or opportunity name")
    parser.add_argument("--list", action="store_true", dest="list_opps",
                        help="List active opportunities")
    parser.add_argument("--all", action="store_true", dest="run_all", default=True,
                        help="Export all active opportunities (default)")
    parser.add_argument("--team", nargs="+", metavar="NAME",
                        help="Filter by multiple tech leads")
    parser.add_argument("-o", "--output-dir", default="output",
                        help="Output directory (default: output/)")
    parser.add_argument("-P", "--prompt-file", metavar="FILE",
                        help="Write output to FILE (single-opp mode)")
    parser.add_argument("-p", "--prompt", action="store_true",
                        help="Print output to stdout (single-opp mode)")
    parser.add_argument("-c", "--cookies", default=None,
                        help="Cookie string")
    parser.add_argument("-f", "--cookies-file", metavar="FILE",
                        help="Read cookies from file")
    parser.add_argument("--type", choices=["json", "txt", "docx"], default="json",
                        help="Output format (default: json)")
    parser.add_argument("--debug", action="store_true",
                        help="Print debug info to stderr")
    args = parser.parse_args()

    cookies = _get_cookies(args)

    if args.evaluation_uuid:
        args.run_all = False

    # --list / --all mode
    if args.list_opps or args.run_all:
        if args.team:
            user_values = _resolve_users(cookies, args.team)
            opps = fetch_opportunities(cookies, user_values=user_values, stage_filter="active")
        else:
            opps = fetch_opportunities(cookies)

        if args.list_opps:
            if args.team:
                print(f"\n{'#':<4} {'Opportunity':<50} {'Tech Lead':<25} {'UUID':<38} {'Stage':<30} {'Deal Size':>12}")
                print("-" * 160)
                for i, o in enumerate(opps, 1):
                    print(f"{i:<4} {o['name']:<50} {o['tech_lead']:<25} {o['uuid']:<38} {o['stage']:<30} ${float(o['deal_size']):>11,.2f}")
            else:
                print(f"\n{'#':<4} {'Opportunity':<55} {'UUID':<38} {'Stage':<30} {'Deal Size':>12}")
                print("-" * 140)
                for i, o in enumerate(opps, 1):
                    print(f"{i:<4} {o['name']:<55} {o['uuid']:<38} {o['stage']:<30} ${float(o['deal_size']):>11,.2f}")
            print(f"\n{len(opps)} opportunities found.\n")
            if not args.run_all:
                return

        os.makedirs(args.output_dir, exist_ok=True)
        print(f"Exporting {len(opps)} opportunities to {args.output_dir}/\n")

        ext = args.type

        def _export_one(opp):
            data = fetch_evaluation_data(opp["uuid"], cookies, debug=args.debug)
            opp_json = build_opportunity_json(data)
            account = opp["name"].split(" - ", 1)[0].strip()
            opp_dir = os.path.join(args.output_dir, _safe_filename(account))
            os.makedirs(opp_dir, exist_ok=True)
            path = os.path.join(opp_dir, f"{_safe_filename(opp['name'])}.{ext}")
            if ext == "docx":
                doc = format_opportunity_docx(opp_json)
                doc.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    if ext == "txt":
                        f.write(format_opportunity_text(opp_json))
                    else:
                        json.dump(build_output_envelope([opp_json]), f, indent=2, ensure_ascii=False)
            return path, opp_json

        exported, skipped = 0, 0
        all_opps = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as pool:
            future_to_opp = {pool.submit(_export_one, opp): opp for opp in opps}
            for future in concurrent.futures.as_completed(future_to_opp):
                opp = future_to_opp[future]
                try:
                    path, opp_json = future.result()
                    exported += 1
                    all_opps.append(opp_json)
                    print(f"  [{exported + skipped}/{len(opps)}] {opp['name']} -> {path}")
                except requests.HTTPError as e:
                    if e.response is not None and e.response.status_code == 403:
                        skipped += 1
                        print(f"  [{exported + skipped}/{len(opps)}] {opp['name']} SKIPPED (no access)")
                    else:
                        print(f"  [{exported + skipped}/{len(opps)}] {opp['name']} FAILED: {e}", file=sys.stderr)
                except requests.RequestException as e:
                    print(f"  {opp['name']} FAILED: {e}", file=sys.stderr)

        if all_opps:
            all_opps.sort(key=lambda o: o.get("Customer", ""))
            combined_ext = "json" if ext == "docx" else ext
            all_path = os.path.join(args.output_dir, f"all.{combined_ext}")
            with open(all_path, "w", encoding="utf-8") as f:
                if combined_ext == "txt":
                    f.write(format_all_opportunities_text(all_opps))
                else:
                    json.dump(build_output_envelope(all_opps), f, indent=2, ensure_ascii=False)
            print(f"  Combined -> {all_path}")

        print(f"\nDone. {exported} exported, {skipped} skipped (no access).")
        return

    # Single-opp mode (argument can be evaluation UUID or opportunity name)
    if not args.evaluation_uuid:
        parser.error("evaluation_uuid or opportunity name required (or use --list / --all)")

    uuid = _resolve_opp_identifier(cookies, args.evaluation_uuid)
    data = fetch_evaluation_data(uuid, cookies, debug=args.debug)
    opp_json = build_opportunity_json(data)

    if args.type == "docx":
        doc = format_opportunity_docx(opp_json)
        out_path = args.prompt_file or f"{_safe_filename(opp_json.get('Opportunity_Name', 'output'))}.docx"
        doc.save(out_path)
        print(f"Written to: {out_path}")
    else:
        if args.type == "txt":
            output = format_opportunity_text(opp_json)
        else:
            output = json.dumps(build_output_envelope([opp_json]), indent=2, ensure_ascii=False)

        if args.prompt_file:
            with open(args.prompt_file, "w", encoding="utf-8") as f:
                f.write(output)
            print(f"Written to: {args.prompt_file}")
        else:
            print(output)


if __name__ == "__main__":
    main()
